#!/usr/bin/env python3
"""
Script unifié pour convertir les fichiers XML, Excel, Word et images en PDF
Inclut l'OCR pour créer des PDF recherchables à partir d'images scannées

Support des formats:
- Images: .jpg, .jpeg, .png, .bmp, .tiff (avec OCR optionnel)
- Excel: .xlsx, .xls, .xlsm, .xlsb
- Word: .docx, .doc
- PowerPoint: .pptx, .ppt
- XML: .xml

Moteurs OCR supportés:
- Tesseract (recommandé pour le français)
- EasyOCR (bon pour documents complexes)
- PaddleOCR (le plus rapide)
"""

import os
import sys
import subprocess
import platform
import xml.dom.minidom
from pathlib import Path
from PIL import Image
import textwrap
import shutil
import time
import csv
from datetime import datetime

# Pour la génération de PDF à partir de XML
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Pour Windows COM (Microsoft Office)
try:
    import win32com.client
    import pythoncom
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False

# Pour la lecture des fichiers Excel (méthode de secours)
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# Pour la lecture des fichiers Word (méthode de secours)
try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# OCR - Tesseract
try:
    import pytesseract
    from pdf2image import convert_from_path
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

# OCR - EasyOCR
try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

# OCR - PaddleOCR
try:
    from paddleocr import PaddleOCR
    PADDLEOCR_AVAILABLE = True
except ImportError:
    PADDLEOCR_AVAILABLE = False

# Pour créer des PDF avec couche de texte
try:
    import PyPDF2
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# Configuration globale
METHODE_CONVERSION = "auto"
REPORTLAB_FALLBACK_ENABLED = False  # Par défaut: pas de fallback ReportLab pour Word/Excel/PPT  # "office", "libreoffice", "reportlab", "auto"
LIBREOFFICE_PATH = None  # Sera détecté automatiquement
UTILISER_OCR = False  # Active l'OCR pour les images
MOTEUR_OCR = "auto"  # "tesseract", "easyocr", "paddleocr", "auto"

# Nom des PDF: conserver l'extension d'origine (ex: x.jpg -> x.jpg.pdf)
KEEP_EXT_IN_NAME = True

# Conversion HTML (.htm/.html) via navigateur headless
BROWSER_PATH = None  # Chrome/Edge détecté automatiquement

# Journalisation (log CSV)
JOURNAL_ENABLED = False
JOURNAL_ERRORS_ONLY = True  # Par défaut: n'écrire dans le journal que les erreurs
JOURNAL_PATH = None
_JOURNAL_FH = None
_JOURNAL_WRITER = None


# Contexte d'erreur (par fichier) pour enrichir le journal CSV
_LAST_ERRORS: list[str] = []
_LAST_INFOS: list[str] = []
_LAST_METHOD_USED: str | None = None
_LAST_EXCEPTION: str | None = None

def reset_error_context():
    """Réinitialise les buffers (erreurs + infos) pour le fichier en cours."""
    global _LAST_ERRORS, _LAST_EXCEPTION, _LAST_INFOS, _LAST_METHOD_USED
    _LAST_ERRORS = []
    _LAST_EXCEPTION = None
    _LAST_INFOS = []
    _LAST_METHOD_USED = None

def log_error(message: str, exc: Exception | str | None = None):
    """Affiche une erreur et la mémorise pour le journal (concaténation).
    - error_messages : concatène les messages
    - exception : stocke la *première* exception détaillée (traceback si possible)
    """
    global _LAST_ERRORS, _LAST_EXCEPTION
    print(message)
    try:
        msg = str(message).strip()
        if msg:
            _LAST_ERRORS.append(msg)
    except Exception:
        pass

    if exc is not None and _LAST_EXCEPTION is None:
        try:
            import traceback as _tb
            if isinstance(exc, BaseException):
                _LAST_EXCEPTION = ''.join(_tb.format_exception(type(exc), exc, exc.__traceback__)).strip()
            else:
                _LAST_EXCEPTION = str(exc).strip()
        except Exception:
            try:
                _LAST_EXCEPTION = str(exc)
            except Exception:
                _LAST_EXCEPTION = None

def init_journal(dossier_base: Path, nom_prefixe: str = "conversion_log"):
    """Initialise un journal CSV dans dossier_base (colonnes étendues)."""
    global JOURNAL_ENABLED, JOURNAL_PATH, _JOURNAL_FH, _JOURNAL_WRITER
    try:
        dossier_base.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        JOURNAL_PATH = dossier_base / f"{nom_prefixe}_{ts}.csv"
        _JOURNAL_FH = open(JOURNAL_PATH, "w", newline="", encoding="utf-8")
        _JOURNAL_WRITER = csv.writer(_JOURNAL_FH)
        _JOURNAL_WRITER.writerow([
            "timestamp",
            "status",
            "filetype",
            "source",
            "dest_pdf",
            "duration_s",
            "details",
            "info_messages",
            "error_messages",
            "exception",
            "method_used",
        ])
        JOURNAL_ENABLED = True
        print(f"🧾 Journal activé: {JOURNAL_PATH}")
    except Exception as e:
        JOURNAL_ENABLED = False
        JOURNAL_PATH = None
        _JOURNAL_FH = None
        _JOURNAL_WRITER = None
        print(f"⚠️  Impossible de créer le journal: {e}")


def journaliser(status: str, source: Path, dest_pdf=None, duration_s=None, details: str = "", error_messages: str = "", exception: Exception | str | None = None, info_messages: str = "", method_used: str | None = None):
    """Écrit une ligne dans le journal si activé.
    Compatibilité: les anciens appels journaliser(status, source, dest_pdf, duration_s, details) continuent à fonctionner.
    """
    global JOURNAL_ENABLED, _JOURNAL_WRITER
    if not JOURNAL_ENABLED or _JOURNAL_WRITER is None:
        return

    # Par défaut: journaliser uniquement les erreurs (failed / skipped_password)
    if JOURNAL_ERRORS_ONLY and status not in ('failed', 'skipped_password'):
        return
    try:
        filetype = source.suffix.lower().lstrip(".")
        _JOURNAL_WRITER.writerow([
            datetime.now().isoformat(timespec="seconds"),
            status,
            filetype,
            str(source),
            str(dest_pdf) if dest_pdf else "",
            f"{duration_s:.3f}" if isinstance(duration_s, (int, float)) else "",
            details or "",
            info_messages or "",
            error_messages or (details or ""),
            str(exception) if exception is not None else "",
            method_used or "",
        ])
    except Exception:
        pass


def fermer_journal():
    """Ferme le journal si ouvert."""
    global _JOURNAL_FH, _JOURNAL_WRITER, JOURNAL_ENABLED
    try:
        if _JOURNAL_FH:
            _JOURNAL_FH.flush()
            _JOURNAL_FH.close()
    finally:
        _JOURNAL_FH = None
        _JOURNAL_WRITER = None
        JOURNAL_ENABLED = False



def detecter_tesseract():
    """Détecte si Tesseract est installé et configuré"""
    if not TESSERACT_AVAILABLE:
        return False
    
    try:
        # Chemins possibles pour Tesseract sur Windows
        chemins_windows = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Users\%USERNAME%\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"
        ]
        
        # Remplacer %USERNAME%
        chemins_windows = [os.path.expandvars(p) for p in chemins_windows]
        
        # Vérifier chaque chemin
        for chemin in chemins_windows:
            if os.path.exists(chemin):
                pytesseract.pytesseract.tesseract_cmd = chemin
                # Tester
                version = pytesseract.get_tesseract_version()
                return True
        
        # Linux/Mac - généralement dans PATH
        try:
            version = pytesseract.get_tesseract_version()
            return True
        except:
            return False
            
    except:
        return False

def ocr_avec_tesseract(chemin_image, langue='fra+eng'):
    """OCR avec Tesseract"""
    if not TESSERACT_AVAILABLE or not detecter_tesseract():
        return None
        
    try:
        image = Image.open(chemin_image)
        
        # Configuration Tesseract pour meilleure qualité
        config = '--oem 3 --psm 6'
        
        # Extraire le texte
        texte = pytesseract.image_to_string(image, lang=langue, config=config)
        
        # Calculer la confiance
        donnees = pytesseract.image_to_data(image, lang=langue, output_type=pytesseract.Output.DICT)
        confidences = [int(conf) for conf in donnees['conf'] if int(conf) > 0]
        confiance = sum(confidences) / len(confidences) if confidences else 0
        
        return {
            'texte': texte,
            'confiance': confiance,
            'moteur': 'tesseract'
        }
    except Exception as e:
        print(f"    ⚠ Erreur Tesseract: {e}")
        return None

def ocr_avec_easyocr(chemin_image):
    """OCR avec EasyOCR"""
    if not EASYOCR_AVAILABLE:
        return None
        
    try:
        # Initialiser le reader (télécharge les modèles au premier usage)
        reader = easyocr.Reader(['fr', 'en'], gpu=False)
        
        # Lire l'image
        resultats = reader.readtext(str(chemin_image))
        
        # Extraire le texte
        texte = '\n'.join([res[1] for res in resultats])
        
        # Calculer la confiance moyenne
        confidences = [res[2] for res in resultats]
        confiance = (sum(confidences) / len(confidences) * 100) if confidences else 0
        
        return {
            'texte': texte,
            'confiance': confiance,
            'moteur': 'easyocr'
        }
    except Exception as e:
        print(f"    ⚠ Erreur EasyOCR: {e}")
        return None

def ocr_avec_paddleocr(chemin_image):
    """OCR avec PaddleOCR"""
    if not PADDLEOCR_AVAILABLE:
        return None
        
    try:
        # Initialiser PaddleOCR
        ocr = PaddleOCR(use_angle_cls=True, lang='french', use_gpu=False)
        
        # Faire l'OCR
        result = ocr.ocr(str(chemin_image), cls=True)
        
        # Extraire le texte
        texte_lignes = []
        confidences = []
        
        if result and result[0]:
            for line in result[0]:
                texte_lignes.append(line[1][0])
                confidences.append(line[1][1])
        
        texte = '\n'.join(texte_lignes)
        confiance = (sum(confidences) / len(confidences) * 100) if confidences else 0
        
        return {
            'texte': texte,
            'confiance': confiance,
            'moteur': 'paddleocr'
        }
    except Exception as e:
        print(f"    ⚠ Erreur PaddleOCR: {e}")
        return None

def choisir_meilleur_ocr(chemin_image):
    """Teste les moteurs OCR disponibles et choisit le meilleur"""
    resultats = []
    
    # Tester Tesseract
    if TESSERACT_AVAILABLE and detecter_tesseract():
        print("    🔤 Test Tesseract...")
        res = ocr_avec_tesseract(chemin_image)
        if res:
            resultats.append(res)
    
    # Tester EasyOCR
    if EASYOCR_AVAILABLE:
        print("    🔤 Test EasyOCR...")
        res = ocr_avec_easyocr(chemin_image)
        if res:
            resultats.append(res)
    
    # Tester PaddleOCR
    if PADDLEOCR_AVAILABLE:
        print("    🔤 Test PaddleOCR...")
        res = ocr_avec_paddleocr(chemin_image)
        if res:
            resultats.append(res)
    
    if not resultats:
        return None
    
    # Choisir le meilleur basé sur la confiance
    meilleur = max(resultats, key=lambda x: x['confiance'])
    print(f"    ✅ Meilleur: {meilleur['moteur']} (confiance: {meilleur['confiance']:.1f}%)")
    
    return meilleur

def creer_pdf_avec_ocr(chemin_image, texte_ocr, chemin_pdf):
    """Crée un PDF avec l'image et le texte OCR pour la recherche"""
    if not REPORTLAB_AVAILABLE:
        return False
    
    try:
        from reportlab.pdfgen import canvas
        
        # Ouvrir l'image
        img = Image.open(chemin_image)
        img_width, img_height = img.size
        
        # Calculer les dimensions pour le PDF
        page_width, page_height = A4
        ratio = min(page_width / img_width, page_height / img_height) * 0.95
        pdf_img_width = img_width * ratio
        pdf_img_height = img_height * ratio
        
        # Centrer l'image
        x_offset = (page_width - pdf_img_width) / 2
        y_offset = (page_height - pdf_img_height) / 2
        
        # Créer le PDF
        c = canvas.Canvas(str(chemin_pdf), pagesize=A4)
        
        # Ajouter l'image
        c.drawImage(str(chemin_image), x_offset, y_offset, 
                   width=pdf_img_width, height=pdf_img_height)
        
        # Ajouter le texte invisible pour la recherche
        c.setFillAlpha(0)  # Texte invisible
        c.setFont("Helvetica", 8)
        
        # Diviser le texte en lignes
        lignes = texte_ocr.split('\n')
        y_position = page_height - 50
        
        for ligne in lignes:
            if ligne.strip():
                # Nettoyer la ligne pour éviter les erreurs
                ligne_propre = ligne.encode('latin-1', 'replace').decode('latin-1')
                c.drawString(50, y_position, ligne_propre)
                y_position -= 12
                if y_position < 50:
                    c.showPage()
                    y_position = page_height - 50
        
        c.save()
        return True
        
    except Exception as e:
        print(f"    ⚠ Erreur création PDF avec OCR: {e}")
        return False

def detecter_libreoffice():
    """Détecte le chemin d'installation de LibreOffice"""
    global LIBREOFFICE_PATH
    
    chemins_possibles = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\Program Files\LibreOffice 7\program\soffice.exe",
        r"C:\Program Files\LibreOffice 24\program\soffice.exe",
    ]
    
    # Ajouter le chemin depuis PATH
    try:
        result = shutil.which("soffice")
        if result:
            chemins_possibles.insert(0, result)
    except:
        pass
    
    for chemin in chemins_possibles:
        if os.path.exists(chemin):
            LIBREOFFICE_PATH = chemin
            return True
    
    return False

def detecter_browser_headless():
    """Détecte Chrome ou Edge pour imprimer du HTML en PDF en mode headless."""
    global BROWSER_PATH

    # Déjà détecté
    if BROWSER_PATH and os.path.exists(BROWSER_PATH):
        return True

    # 1) PATH
    for exe in ("chrome", "chrome.exe", "msedge", "msedge.exe"):
        p = shutil.which(exe)
        if p and os.path.exists(p):
            BROWSER_PATH = p
            return True

    # 2) Chemins Windows usuels
    chemins_windows = [
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    ]
    for p in chemins_windows:
        if os.path.exists(p):
            BROWSER_PATH = p
            return True

    # 3) macOS / Linux (si jamais)
    for exe in ("google-chrome", "chromium", "chromium-browser", "microsoft-edge"):
        p = shutil.which(exe)
        if p and os.path.exists(p):
            BROWSER_PATH = p
            return True

    return False

def convertir_html_vers_pdf(chemin_source, chemin_pdf):
    """
    Convertit un fichier HTML en PDF avec Chrome/Edge headless (rendu fidèle).
    Nécessite un navigateur Chromium.
    """
    if not detecter_browser_headless():
        print("  ❌ Aucun navigateur (Chrome/Edge) détecté pour convertir le HTML.")
        return False

    try:
        # file:/// URI
        source_uri = Path(chemin_source).absolute().as_uri()

        # Chrome/Edge écrivent dans un répertoire; on donne un chemin complet.
        # On utilise un profil temporaire pour éviter les conflits.
        tmp_profile = Path(chemin_pdf).parent / f".tmp_profile_{int(time.time()*1000)}"
        tmp_profile.mkdir(parents=True, exist_ok=True)

        cmd = [
            BROWSER_PATH,
            "--headless=new",
            "--disable-gpu",
            "--no-first-run",
            "--no-default-browser-check",
            f"--user-data-dir={str(tmp_profile)}",
            f"--print-to-pdf={str(Path(chemin_pdf).absolute())}",
            "--print-to-pdf-no-header",
            source_uri,
        ]

        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

        # Nettoyage profil
        try:
            shutil.rmtree(tmp_profile, ignore_errors=True)
        except:
            pass

        if result.returncode != 0:
            log_error(f"  ⚠ Erreur navigateur: {result.stderr.strip()[:300]}", result.stderr)
            return False

        return Path(chemin_pdf).exists() and Path(chemin_pdf).stat().st_size > 0

    except Exception as e:
        log_error(f"  ⚠ Erreur HTML->PDF: {e}", e)
        return False

def convertir_texte_vers_pdf(chemin_source, chemin_pdf, titre=None):
    """Convertit un fichier texte (.txt/.log) en PDF propre (monospace, pagination)."""
    if not REPORTLAB_AVAILABLE:
        print("  ❌ ReportLab requis pour convertir les fichiers texte.")
        return False

    try:
        contenu = Path(chemin_source).read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        print(f"  ⚠ Erreur lecture texte: {e}")
        return False

    try:
        styles = getSampleStyleSheet()
        mono = ParagraphStyle(
            "Mono",
            parent=styles["Normal"],
            fontName="Courier",
            fontSize=9,
            leading=11,
        )

        doc = SimpleDocTemplate(str(chemin_pdf), pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
        story = []

        if titre:
            story.append(Paragraph(titre, styles["Heading2"]))
            story.append(Spacer(1, 12))

        # Preformatted conserve les retours à la ligne et l'indentation
        story.append(Preformatted(contenu, mono))
        doc.build(story)
        return True

    except Exception as e:
        print(f"  ⚠ Erreur ReportLab texte: {e}")
        return False

def convertir_msg_vers_pdf(chemin_source, chemin_pdf):
    """
    Convertit un .msg en PDF.
    - Windows recommandé: Outlook via COM -> export HTML -> HTML->PDF (Chrome/Edge headless)
    - Fallback: extraction basique si possible.
    """
    extension = Path(chemin_source).suffix.lower()
    if extension != ".msg":
        return False

    # 1) Outlook COM si possible
    if WIN32COM_AVAILABLE and platform.system() == "Windows":
        try:
            pythoncom.CoInitialize()
            outlook = win32com.client.Dispatch("Outlook.Application")
            session = outlook.Session

            # OpenSharedItem est le plus fiable pour .msg
            item = None
            try:
                item = session.OpenSharedItem(str(Path(chemin_source).absolute()))
            except Exception:
                try:
                    item = outlook.CreateItemFromTemplate(str(Path(chemin_source).absolute()))
                except Exception:
                    item = None

            if item is None:
                raise RuntimeError("Impossible d'ouvrir le .msg via Outlook.")

            tmp_html = Path(chemin_pdf).with_suffix(".tmp.html")
            # olHTML = 5
            item.SaveAs(str(tmp_html), 5)

            # Convertir HTML -> PDF
            ok = convertir_html_vers_pdf(tmp_html, chemin_pdf)

            try:
                tmp_html.unlink(missing_ok=True)
            except Exception:
                pass

            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

            return ok

        except Exception as e:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
            log_error(f"  ⚠ Erreur MSG via Outlook: {e}", e)
            # On tombe sur fallback

    # 2) Fallback: extraire en texte simple
    try:
        # Tentative d'utiliser extract_msg si installé
        try:
            import extract_msg  # type: ignore
            msg = extract_msg.Message(str(chemin_source))
            msg.process()
            texte = f"Subject: {msg.subject}\nFrom: {msg.sender}\nTo: {msg.to}\nDate: {msg.date}\n\n{msg.body or ''}"
            tmp_txt = Path(chemin_pdf).with_suffix(".tmp.txt")
            tmp_txt.write_text(texte, encoding="utf-8", errors="replace")
            ok = convertir_texte_vers_pdf(tmp_txt, chemin_pdf, titre=Path(chemin_source).name)
            try:
                tmp_txt.unlink(missing_ok=True)
            except Exception:
                pass
            return ok
        except ImportError:
            print("  ❌ Ni Outlook COM ni extract_msg disponible pour convertir .msg.")
            return False

    except Exception as e:
        log_error(f"  ⚠ Erreur MSG fallback: {e}", e)
        return False


def log_info(message: str):
    """Affiche une info et la mémorise pour le journal (concaténation)."""
    global _LAST_INFOS
    print(message)
    try:
        msg = str(message).strip()
        if msg:
            _LAST_INFOS.append(msg)
    except Exception:
        pass


def is_password_error(err: Exception | str) -> bool:
    """Heuristique: détecte si une erreur indique un fichier protégé par mot de passe.
    Objectif: SKIP propre (pas de fallback) pour éviter des PDF incomplets.
    """
    try:
        msg = str(err).lower()
    except Exception:
        return False
    keywords = [
        "password", "mot de passe", "mdp",
        "protected", "protég", "protection", "encrypt", "encrypted", "chiffr",
        "cannot be opened because it is password", "the password is incorrect",
        "requires a password", "un mot de passe est requis",
    ]
    return any(k in msg for k in keywords)

def detecter_office():
    """Vérifie si Microsoft Office est installé et accessible via COM"""
    if not WIN32COM_AVAILABLE:
        return False
    
    try:
        # Tester Word
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Quit()
        
        # Tester Excel
        excel = win32com.client.Dispatch("Excel.Application")
        excel.Quit()
        
        pythoncom.CoUninitialize()
        return True
    except:
        try:
            pythoncom.CoUninitialize()
        except:
            pass
        return False

def convertir_avec_office(chemin_source, chemin_pdf):
    """Convertit un fichier Word/Excel/PowerPoint en PDF en utilisant Microsoft Office"""
    if not WIN32COM_AVAILABLE:
        return False
    
    extension = chemin_source.suffix.lower()
    _pw = is_password_error

    
    try:
        pythoncom.CoInitialize()
        
        # Chemins absolus nécessaires pour COM
        chemin_source_abs = str(chemin_source.absolute())
        chemin_pdf_abs = str(chemin_pdf.absolute())
        
        if extension in ['.doc', '.docx']:
            # Conversion Word
            # IMPORTANT: utiliser DispatchEx pour créer une nouvelle instance Word
            # (évite les problèmes quand Word est déjà ouvert et qu'on se connecte à une instance existante)
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            # 0 = wdAlertsNone (plus fiable que False selon les versions)
            word.DisplayAlerts = 0
            # 3 = msoAutomationSecurityForceDisable (désactive les macros pendant l'automatisation)
            try:
                word.AutomationSecurity = 3
            except Exception:
                pass
            
            doc = None
            try:
                # Ouverture en lecture seule, sans ajouts "fichiers récents", sans dialogues
                doc = word.Documents.Open(
                    chemin_source_abs,
                    ReadOnly=True,
                    AddToRecentFiles=False,
                    ConfirmConversions=False,
                    NoEncodingDialog=True
                )
                
                # Export PDF : ExportAsFixedFormat est généralement le plus robuste
                # 17 = wdExportFormatPDF
                try:
                    doc.ExportAsFixedFormat(
                        OutputFileName=chemin_pdf_abs,
                        ExportFormat=17,
                        OpenAfterExport=False,
                        OptimizeFor=0,      # 0 = wdExportOptimizeForPrint
                        Range=0,            # 0 = wdExportAllDocument
                        Item=0,             # 0 = wdExportDocumentContent
                        IncludeDocProps=True,
                        KeepIRM=True,
                        CreateBookmarks=1,  # 1 = wdExportCreateHeadingBookmarks
                        DocStructureTags=True,
                        BitmapMissingFonts=True,
                        UseISO19005_1=False
                    )
                    succes = True
                except Exception as e_export:
                    # Fallback : SaveAs2/FileFormat=17 (marche parfois quand ExportAsFixedFormat échoue)
                    try:
                        if hasattr(doc, "SaveAs2"):
                            doc.SaveAs2(chemin_pdf_abs, FileFormat=17)
                        else:
                            doc.SaveAs(chemin_pdf_abs, FileFormat=17)
                        succes = True
                    except Exception:
                        if _pw(e_export):
                            log_error(f"  🔒 Document protégé par mot de passe (Word) : {e_export}", e_export)
                            succes = "password"
                        else:
                            log_error(f"  ⚠ Erreur Office Word: {e_export}", e_export)
                            succes = False
            except Exception as e:
                if _pw(e):
                    log_error(f"  🔒 Document protégé par mot de passe (Word) : {e}", e)
                    succes = "password"
                else:
                    log_error(f"  ⚠ Erreur Office Word: {e}", e)
                    succes = False
            finally:
                try:
                    if doc is not None:
                        doc.Close(False)
                except Exception:
                    pass
                try:
                    word.Quit()
                except Exception:
                    pass

        elif extension in ['.xls', '.xlsx', '.xlsm', '.xlsb']:
            # Conversion Excel
            excel = win32com.client.Dispatch("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            
            try:
                wb = excel.Workbooks.Open(chemin_source_abs, ReadOnly=True, UpdateLinks=0)
                # Type PDF = 0
                wb.ExportAsFixedFormat(0, chemin_pdf_abs)
                wb.Close(False)
                succes = True
            except Exception as e:
                if _pw(e):
                    log_error(f"  🔒 Fichier protégé par mot de passe (Excel) : {e}", e)
                    succes = "password"
                else:
                    log_error(f"  ⚠ Erreur Office Excel: {e}", e)
                    succes = False
            finally:
                excel.Quit()
        
        elif extension in ['.ppt', '.pptx']:
            # Conversion PowerPoint
            powerpoint = win32com.client.Dispatch("PowerPoint.Application")
            powerpoint.Visible = False
            powerpoint.DisplayAlerts = False
            
            try:
                presentation = powerpoint.Presentations.Open(chemin_source_abs, WithWindow=False)
                # Format PDF = 32
                presentation.SaveAs(chemin_pdf_abs, 32)
                presentation.Close()
                succes = True
            except Exception as e:
                log_error(f"  ⚠ Erreur Office PowerPoint: {e}", e)
                succes = False
            finally:
                powerpoint.Quit()
        
        else:
            succes = False
        
        pythoncom.CoUninitialize()
        return succes
        
    except Exception as e:
        log_error(f"  ⚠ Erreur COM générale: {e}", e)
        try:
            pythoncom.CoUninitialize()
        except:
            pass
        return False

def convertir_avec_libreoffice(chemin_source, chemin_pdf):
    """Convertit un fichier en PDF en utilisant LibreOffice en mode headless"""
    if not LIBREOFFICE_PATH:
        return False
    
    try:
        # LibreOffice nécessite le répertoire de sortie, pas le fichier
        repertoire_sortie = chemin_pdf.parent
        
        # Commande LibreOffice avec options d'encodage
        cmd = [
            LIBREOFFICE_PATH,
            '--headless',
            '--convert-to', 'pdf:writer_pdf_Export',
            '--infilter=UTF8',  # Forcer l'encodage UTF-8
            '--outdir', str(repertoire_sortie),
            str(chemin_source)
        ]
        
        # Variables d'environnement pour l'encodage
        env = os.environ.copy()
        env['PYTHONIOENCODING'] = 'utf-8'
        
        # Exécuter la conversion
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60, env=env)
        
        if result.returncode == 0:
            # LibreOffice crée le PDF avec le même nom que le source
            pdf_genere = repertoire_sortie / (chemin_source.stem + '.pdf')
            
            # Renommer si nécessaire
            if pdf_genere != chemin_pdf:
                if chemin_pdf.exists():
                    chemin_pdf.unlink()
                pdf_genere.rename(chemin_pdf)
            
            return True
        else:
            log_error(f"  ⚠ Erreur LibreOffice: {result.stderr}", result.stderr)
            return False
            
    except subprocess.TimeoutExpired:
        log_error("  ⚠ Timeout LibreOffice (>60s)")
        return False
    except Exception as e:
        log_error(f"  ⚠ Erreur LibreOffice: {e}", e)
        return False

def convertir_jpg_vers_pdf(chemin_jpg, chemin_pdf):
    """Convertit une image en PDF avec option OCR"""
    # Si OCR activé
    if UTILISER_OCR:
        print(f"  🔤 OCR activé pour: {chemin_jpg.name}")
        
        # Déterminer quel moteur OCR utiliser
        resultat_ocr = None
        
        if MOTEUR_OCR == "auto":
            # Choisir automatiquement le meilleur
            resultat_ocr = choisir_meilleur_ocr(chemin_jpg)
        elif MOTEUR_OCR == "tesseract":
            resultat_ocr = ocr_avec_tesseract(chemin_jpg)
        elif MOTEUR_OCR == "easyocr":
            resultat_ocr = ocr_avec_easyocr(chemin_jpg)
        elif MOTEUR_OCR == "paddleocr":
            resultat_ocr = ocr_avec_paddleocr(chemin_jpg)
        
        if resultat_ocr and resultat_ocr['texte']:
            # Afficher un extrait du texte reconnu
            texte = resultat_ocr['texte']
            print(f"    📝 Texte extrait ({len(texte)} caractères)")
            if len(texte) > 100:
                print(f"    → {texte[:100]}...")
            
            # Créer le PDF avec le texte OCR
            if creer_pdf_avec_ocr(chemin_jpg, texte, chemin_pdf):
                return True
            else:
                print("    ⚠ Échec création PDF avec OCR, conversion standard...")
        else:
            print("    ⚠ OCR échoué, conversion standard...")
    
    # Conversion standard sans OCR
    try:
        with Image.open(chemin_jpg) as img:
            if img.mode != 'RGB':
                img = img.convert('RGB')
            img.save(chemin_pdf, "PDF", resolution=100.0, quality=95)
        return True
    except Exception as e:
        log_error(f"  ⚠ Erreur conversion image: {e}", e)
        return False

def convertir_xml_vers_pdf(chemin_xml, chemin_pdf):
    """Convertit un fichier XML en PDF avec formatage"""
    if not REPORTLAB_AVAILABLE:
        print("  ⚠ ReportLab non installé pour XML")
        return False
    
    try:
        # Lire et formater le XML
        with open(chemin_xml, 'r', encoding='utf-8') as f:
            contenu_xml = f.read()
        
        # Parser pour un joli formatage
        try:
            dom = xml.dom.minidom.parseString(contenu_xml)
            xml_formate = dom.toprettyxml(indent="  ")
            lignes = [ligne for ligne in xml_formate.split('\n') if ligne.strip()]
            xml_formate = '\n'.join(lignes)
        except:
            xml_formate = contenu_xml
        
        # Créer le PDF
        doc = SimpleDocTemplate(str(chemin_pdf), pagesize=A4)
        styles = getSampleStyleSheet()
        
        style_titre = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            textColor=colors.darkblue
        )
        
        style_code = ParagraphStyle(
            'CodeStyle',
            parent=styles['Code'],
            fontSize=8,
            leftIndent=20,
            fontName='Courier'
        )
        
        story = []
        
        # Titre
        titre = f"Fichier XML: {chemin_xml.name}"
        story.append(Paragraph(titre, style_titre))
        story.append(Spacer(1, 12))
        
        # Contenu
        lignes_xml = xml_formate.split('\n')
        for ligne in lignes_xml:
            if ligne.strip():
                ligne_echappee = ligne.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Preformatted(ligne_echappee, style_code))
        
        doc.build(story)
        return True
        
    except Exception as e:
        log_error(f"  ⚠ Erreur conversion XML: {e}", e)
        return False

def convertir_word_vers_pdf_reportlab(chemin_word, chemin_pdf):
    """Méthode de secours pour Word avec ReportLab - Version améliorée"""
    if not REPORTLAB_AVAILABLE or not PYTHON_DOCX_AVAILABLE:
        return False
    
    if chemin_word.suffix.lower() != '.docx':
        print("  ⚠ ReportLab ne supporte que .docx")
        return False
    
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.units import cm
        
        # Charger une police qui supporte les accents
        try:
            pdfmetrics.registerFont(TTFont('Arial', 'C:/Windows/Fonts/arial.ttf'))
            font_name = 'Arial'
        except:
            font_name = 'Helvetica'
        
        doc = Document(chemin_word)
        
        # Configuration avec marges plus larges pour Word
        doc_pdf = SimpleDocTemplate(
            str(chemin_pdf), 
            pagesize=A4,
            leftMargin=2.5*cm,
            rightMargin=2.5*cm,
            topMargin=2.5*cm,
            bottomMargin=2.5*cm
        )
        
        styles = getSampleStyleSheet()
        
        # Styles personnalisés
        style_normal = ParagraphStyle(
            'WordNormal',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=11,
            leading=14,
            spaceAfter=6
        )
        
        style_heading1 = ParagraphStyle(
            'WordHeading1',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=16,
            spaceAfter=12,
            spaceBefore=12,
            textColor=colors.HexColor('#2F5597')
        )
        
        style_heading2 = ParagraphStyle(
            'WordHeading2',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=14,
            spaceAfter=10,
            spaceBefore=10,
            textColor=colors.HexColor('#2F5597')
        )
        
        story = []
        
        # Titre du document
        titre = chemin_word.stem.replace('_', ' ')
        story.append(Paragraph(titre, style_heading1))
        story.append(Spacer(1, 20))
        
        # Parcourir les paragraphes et tableaux dans l'ordre
        for element in doc.element.body:
            if element.tag.endswith('p'):
                # Traiter les paragraphes
                for para in doc.paragraphs:
                    if para._element == element:
                        texte = para.text.strip()
                        if texte:
                            # Déterminer le style
                            if para.style.name.startswith('Heading 1'):
                                style = style_heading1
                            elif para.style.name.startswith('Heading 2'):
                                style = style_heading2
                            else:
                                style = style_normal
                            
                            # Gérer les caractères spéciaux
                            texte_escape = (texte
                                .replace('&', '&amp;')
                                .replace('<', '&lt;')
                                .replace('>', '&gt;'))
                            
                            story.append(Paragraph(texte_escape, style))
                        break
            
            elif element.tag.endswith('tbl'):
                # Traiter les tableaux
                for table in doc.tables:
                    if table._element == element:
                        donnees_tableau = []
                        
                        for row in table.rows:
                            ligne = []
                            for cell in row.cells:
                                texte = cell.text.strip()
                                # Limiter la longueur
                                if len(texte) > 100:
                                    texte = texte[:97] + "..."
                                ligne.append(texte)
                            donnees_tableau.append(ligne)
                        
                        if donnees_tableau:
                            # Calculer les largeurs
                            nb_cols = len(donnees_tableau[0])
                            largeur_dispo = A4[0] - 5*cm
                            largeur_col = largeur_dispo / nb_cols
                            
                            t = Table(donnees_tableau, colWidths=[largeur_col]*nb_cols)
                            
                            # Style du tableau Word
                            t.setStyle(TableStyle([
                                ('FONTNAME', (0, 0), (-1, -1), font_name),
                                ('FONTSIZE', (0, 0), (-1, -1), 9),
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F2F2')]),
                                ('TOPPADDING', (0, 0), (-1, -1), 4),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                            ]))
                            
                            story.append(t)
                            story.append(Spacer(1, 12))
                        break
        
        # Note de conversion
        note_style = ParagraphStyle(
            'Note',
            parent=styles['Italic'],
            fontSize=8,
            textColor=colors.grey,
            alignment=1  # Centré
        )
        story.append(Spacer(1, 30))
        story.append(Paragraph(
            "Document converti avec ReportLab - Pour une meilleure qualité, utilisez Microsoft Office ou LibreOffice",
            note_style
        ))
        
        doc_pdf.build(story)
        return True
        
    except Exception as e:
        print(f"  ⚠ Erreur ReportLab Word: {e}")
        import traceback
        traceback.print_exc()
        return False

def convertir_excel_vers_pdf_reportlab(chemin_excel, chemin_pdf):
    """Méthode de secours pour Excel avec ReportLab - Version améliorée"""
    if not REPORTLAB_AVAILABLE or not PANDAS_AVAILABLE:
        return False
    
    try:
        # Import supplémentaire pour meilleure gestion des caractères
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
        
        # Essayer de charger une police qui supporte mieux les accents
        try:
            # Windows
            pdfmetrics.registerFont(TTFont('Arial', 'C:/Windows/Fonts/arial.ttf'))
            font_name = 'Arial'
        except:
            try:
                # Alternative
                pdfmetrics.registerFont(TTFont('DejaVu', 'C:/Windows/Fonts/DejaVuSans.ttf'))
                font_name = 'DejaVu'
            except:
                font_name = 'Helvetica'  # Police par défaut
        
        # Lire TOUTES les feuilles pour une conversion complète
        excel_file = pd.ExcelFile(chemin_excel, engine='openpyxl' if OPENPYXL_AVAILABLE else None)
        
        # Configuration du document avec marges réduites pour plus d'espace
        doc = SimpleDocTemplate(
            str(chemin_pdf), 
            pagesize=A4,
            leftMargin=20,
            rightMargin=20,
            topMargin=30,
            bottomMargin=30
        )
        
        styles = getSampleStyleSheet()
        story = []
        
        # Style personnalisé pour le titre avec meilleure police
        style_titre = ParagraphStyle(
            'ExcelTitle',
            parent=styles['Title'],
            fontName='Helvetica',
            fontSize=14,
            spaceAfter=15
        )
        
        # Titre principal
        titre_text = chemin_excel.name
        story.append(Paragraph(titre_text, style_titre))
        story.append(Spacer(1, 10))
        
        # Traiter chaque feuille
        for sheet_name in excel_file.sheet_names:
            # Lire la feuille avec plus d'options pour préserver le formatage
            df = pd.read_excel(
                chemin_excel, 
                sheet_name=sheet_name, 
                engine='openpyxl' if OPENPYXL_AVAILABLE else None,
                na_filter=True,
                keep_default_na=True
            )
            
            # Ignorer les feuilles vides
            if df.empty:
                continue
            
            # Titre de la feuille
            story.append(Paragraph(f"Feuille: {sheet_name}", styles['Heading2']))
            story.append(Spacer(1, 10))
            
            # Préparer les données avec meilleure gestion des valeurs nulles
            donnees = []
            
            # En-têtes avec gestion des caractères spéciaux
            headers = []
            for col in df.columns:
                header = str(col)
                # Nettoyer les en-têtes "Unnamed"
                if 'Unnamed:' in header:
                    header = ''
                headers.append(header)
            donnees.append(headers)
            
            # Limiter le nombre de lignes mais augmenter la limite
            max_rows = 100
            df_display = df.head(max_rows)
            
            # Ajouter les données avec meilleure gestion des caractères
            for idx, row in df_display.iterrows():
                ligne = []
                for val in row:
                    if pd.isna(val) or str(val) == 'nan':
                        cell_text = ''
                    else:
                        cell_text = str(val)
                        # Limiter la longueur mais augmenter la limite
                        if len(cell_text) > 50:
                            cell_text = cell_text[:47] + '...'
                    ligne.append(cell_text)
                donnees.append(ligne)
            
            if len(donnees) > 1:
                # Calculer automatiquement les largeurs de colonnes
                num_cols = len(donnees[0])
                page_width = A4[0] - 40  # Largeur disponible
                
                # Analyser le contenu pour déterminer les largeurs
                col_widths = []
                for i in range(num_cols):
                    max_len = max(len(str(row[i])) for row in donnees[:20])  # Échantillon
                    col_widths.append(max_len)
                
                # Normaliser les largeurs
                total_width = sum(col_widths)
                if total_width > 0:
                    col_widths = [(w/total_width) * page_width for w in col_widths]
                else:
                    col_widths = [page_width/num_cols] * num_cols
                
                # Largeur minimale et maximale
                min_width = 30
                max_width = 150
                col_widths = [max(min_width, min(max_width, w)) for w in col_widths]
                
                # Créer le tableau
                table = Table(donnees, colWidths=col_widths, repeatRows=1)
                
                # Style amélioré du tableau
                table_style = TableStyle([
                    # En-têtes
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('FONTNAME', (0, 0), (-1, 0), font_name + '-Bold' if font_name != 'Helvetica' else 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('TOPPADDING', (0, 0), (-1, 0), 8),
                    
                    # Données
                    ('FONTNAME', (0, 1), (-1, -1), font_name),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 1), (-1, -1), 'TOP'),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    
                    # Grille
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('LINEBELOW', (0, 0), (-1, 0), 1, colors.HexColor('#2F5597')),
                    
                    # Alternance de couleurs pour les lignes
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F2F2')]),
                    
                    # Padding
                    ('LEFTPADDING', (0, 0), (-1, -1), 4),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                    ('TOPPADDING', (0, 1), (-1, -1), 3),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 3),
                ])
                
                table.setStyle(table_style)
                story.append(table)
                
                # Info sur la limitation
                if len(df) > max_rows:
                    info_style = ParagraphStyle(
                        'InfoStyle',
                        parent=styles['Italic'],
                        fontSize=8,
                        textColor=colors.grey
                    )
                    info_text = f"* Affichage limité aux {max_rows} premières lignes sur {len(df)} au total"
                    story.append(Spacer(1, 5))
                    story.append(Paragraph(info_text, info_style))
            
            story.append(Spacer(1, 20))
        
        # Générer le PDF
        doc.build(story)
        return True
        
    except Exception as e:
        print(f"  ⚠ Erreur ReportLab Excel: {e}")
        import traceback
        traceback.print_exc()
        return False

def convertir_fichier_intelligent(chemin_source, chemin_pdf, methode_forcee=None):
    """Convertit un fichier en utilisant la meilleure méthode disponible"""
    extension = chemin_source.suffix.lower()
    methode = methode_forcee or METHODE_CONVERSION
    
    # Images - toujours avec PIL
    if extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.webp']:
        return convertir_jpg_vers_pdf(chemin_source, chemin_pdf)
    
    # XML - toujours avec ReportLab
    if extension == '.xml':
        return convertir_xml_vers_pdf(chemin_source, chemin_pdf)


    # PDF déjà PDF : si on écrit dans un autre répertoire, copier ; sinon ignorer
    if extension == '.pdf':
        try:
            if Path(chemin_source).absolute() == Path(chemin_pdf).absolute():
                # même fichier (no-op)
                return True
            shutil.copy2(str(chemin_source), str(chemin_pdf))
            return True
        except Exception as e:
            log_error(f"  ⚠ Erreur copie PDF: {e}", e)
            return False

    # Texte brut
    if extension in ['.txt', '.log']:
        titre = f"{chemin_source.name}"
        return convertir_texte_vers_pdf(chemin_source, chemin_pdf, titre=titre)

    # HTML
    if extension in ['.htm', '.html']:
        return convertir_html_vers_pdf(chemin_source, chemin_pdf)

    # Outlook MSG
    if extension == '.msg':
        return convertir_msg_vers_pdf(chemin_source, chemin_pdf)
    
    # Fichiers Office
    if extension in ['.doc', '.docx', '.rtf', '.odt', '.xls', '.xlsx', '.xlsm', '.xlsb', '.ppt', '.pptx']:
        # Ordre de préférence des méthodes
        if methode == "auto":
            # 1. Essayer Microsoft Office
            if WIN32COM_AVAILABLE and detecter_office():
                log_info("  🔧 Tentative Microsoft Office (COM)...")
                print(f"  🔧 Utilisation de Microsoft Office...")
                res_office = convertir_avec_office(chemin_source, chemin_pdf)
                if res_office == "password":
                    _LAST_METHOD_USED = 'office'
                    return "skipped_password"
                if res_office is True:
                    _LAST_METHOD_USED = 'office'
                    return True
                else:
                    log_error("  ❌ Échec Microsoft Office (COM)")
            
            else:
                if not WIN32COM_AVAILABLE:
                    log_info("  ℹ️  Microsoft Office COM indisponible (pywin32 non installé)")
                else:
                    log_info("  ℹ️  Microsoft Office non détecté ou inaccessible via COM")

            # 2. Essayer LibreOffice
            if LIBREOFFICE_PATH:
                log_info(f"  🔧 Tentative LibreOffice... ({LIBREOFFICE_PATH})")
                print(f"  🔧 Utilisation de LibreOffice...")
                if convertir_avec_libreoffice(chemin_source, chemin_pdf):
                    _LAST_METHOD_USED = 'libreoffice'
                    return True
                else:
                    log_error("  ❌ Échec LibreOffice")
            
            else:
                log_info("  ℹ️  LibreOffice non détecté")

            # 3. Méthode de secours ReportLab
            log_info("  🔧 Tentative ReportLab (qualité réduite)...")
            if extension in ['.docx']:
                _LAST_METHOD_USED = 'reportlab'
                ok = convertir_word_vers_pdf_reportlab(chemin_source, chemin_pdf)
                if not ok:
                    log_error("  ❌ Échec ReportLab (Word)")
                return ok
            elif extension in ['.xls', '.xlsx', '.xlsm']:
                _LAST_METHOD_USED = 'reportlab'
                ok = convertir_excel_vers_pdf_reportlab(chemin_source, chemin_pdf)
                if not ok:
                    log_error("  ❌ Échec ReportLab (Excel)")
                return ok
            else:
                log_error(f"  ❌ Aucun fallback ReportLab pour {extension}")
                return False
        
        elif methode == "office":
            res_office = convertir_avec_office(chemin_source, chemin_pdf)
            if res_office == "password":
                return "skipped_password"
            return bool(res_office)
        
        elif methode == "libreoffice":
            return convertir_avec_libreoffice(chemin_source, chemin_pdf)
        
        elif methode == "reportlab":
            if extension in ['.docx']:
                _LAST_METHOD_USED = 'reportlab'
                ok = convertir_word_vers_pdf_reportlab(chemin_source, chemin_pdf)
                if not ok:
                    log_error("  ❌ Échec ReportLab (Word)")
                return ok
            elif extension in ['.xls', '.xlsx', '.xlsm']:
                _LAST_METHOD_USED = 'reportlab'
                ok = convertir_excel_vers_pdf_reportlab(chemin_source, chemin_pdf)
                if not ok:
                    log_error("  ❌ Échec ReportLab (Excel)")
                return ok
            else:
                log_error(f"  ❌ Aucun fallback ReportLab pour {extension}")
                return False
    
    return False

def afficher_configuration():
    """Affiche la configuration détectée"""
    print("\n=== CONFIGURATION DÉTECTÉE ===")
    print(f"Système: {platform.system()} {platform.release()}")
    
    print("\nMéthodes de conversion disponibles:")
    
    # Microsoft Office
    if WIN32COM_AVAILABLE:
        if detecter_office():
            print("✅ Microsoft Office (COM) - DISPONIBLE")
        else:
            print("❌ Microsoft Office (COM) - pywin32 installé mais Office non détecté")
    else:
        print("❌ Microsoft Office (COM) - pywin32 non installé")
    
    # LibreOffice
    if detecter_libreoffice():
        print(f"✅ LibreOffice - DISPONIBLE ({LIBREOFFICE_PATH})")
    else:
        print("❌ LibreOffice - Non détecté")
    
    # ReportLab
    if REPORTLAB_AVAILABLE:
        print("✅ ReportLab (méthode basique) - DISPONIBLE")
        if PANDAS_AVAILABLE:
            print("  ✅ Pandas (pour Excel)")
        else:
            print("  ❌ Pandas (pour Excel)")
        if PYTHON_DOCX_AVAILABLE:
            print("  ✅ python-docx (pour Word)")
        else:
            print("  ❌ python-docx (pour Word)")
    else:
        print("❌ ReportLab - Non installé")
    
    # OCR
    print("\nMoteurs OCR disponibles:")
    ocr_disponibles = []
    
    if TESSERACT_AVAILABLE and detecter_tesseract():
        print("✅ Tesseract OCR - DISPONIBLE")
        try:
            langues = pytesseract.get_languages()
            if 'fra' in langues:
                print("  ✅ Langue française installée")
            else:
                print("  ❌ Langue française non installée")
        except:
            pass
        ocr_disponibles.append('tesseract')
    else:
        print("❌ Tesseract OCR - Non disponible")
        if TESSERACT_AVAILABLE:
            print("  → pytesseract installé mais Tesseract.exe non trouvé")
    
    if EASYOCR_AVAILABLE:
        print("✅ EasyOCR - DISPONIBLE")
        print("  → Téléchargera les modèles au premier usage")
        ocr_disponibles.append('easyocr')
    else:
        print("❌ EasyOCR - Non installé (pip install easyocr)")
    
    if PADDLEOCR_AVAILABLE:
        print("✅ PaddleOCR - DISPONIBLE")
        ocr_disponibles.append('paddleocr')
    else:
        print("❌ PaddleOCR - Non installé (pip install paddlepaddle paddleocr)")
    
    if not ocr_disponibles:
        print("\n⚠️  Aucun moteur OCR disponible")
        print("   Installation recommandée:")
        print("   • Tesseract: https://github.com/UB-Mannheim/tesseract/wiki")
        print("   • pip install pytesseract easyocr")
    
    print(f"\nMéthode de conversion actuelle: {METHODE_CONVERSION}")
    
    print("================================\n")

def convertir_fichier(chemin_source, repertoire_sortie=None, conserver_original=True, forcer=False):
    """Convertit un fichier en PDF"""
    extension = chemin_source.suffix.lower()
    
    # Déterminer le répertoire de sortie
    if repertoire_sortie is None:
        repertoire_dest = chemin_source.parent
    else:
        repertoire_dest = repertoire_sortie
        repertoire_dest.mkdir(parents=True, exist_ok=True)
    
    # Créer le nom du fichier PDF
    if KEEP_EXT_IN_NAME:
        nom_pdf = chemin_source.name + '.pdf'
    else:
        nom_pdf = chemin_source.stem + '.pdf'
    chemin_pdf = repertoire_dest / nom_pdf

    # Cas particulier: source déjà en PDF
    # - si on sort dans le même dossier, on ignore (évite no-op et suppression accidentelle)
    # - si un dossier de sortie est fourni, on copiera le PDF (dans convertir_fichier_intelligent)
    if extension == '.pdf' and repertoire_dest == chemin_source.parent:
        print(f"⏭️  Ignoré (déjà PDF): {chemin_source.name}")
        journaliser('skipped_pdf', chemin_source, None, None, 'déjà PDF (même dossier)', error_messages=' | '.join(_LAST_ERRORS), exception=_LAST_EXCEPTION, info_messages=' | '.join(_LAST_INFOS), method_used=_LAST_METHOD_USED)
        return 'skipped'
    
    # Vérifier si le fichier PDF existe déjà
    if chemin_pdf.exists() and not forcer:
        print(f"⏭️  Ignoré (PDF existant): {chemin_source.name}")
        journaliser('skipped_exists', chemin_source, chemin_pdf, None, 'PDF existant', error_messages=' | '.join(_LAST_ERRORS), exception=_LAST_EXCEPTION, info_messages=' | '.join(_LAST_INFOS), method_used=_LAST_METHOD_USED)
        return 'skipped'
    
    # Gérer les conflits de noms
    if chemin_pdf.exists() and forcer:
        print(f"🔄 Remplacement: {chemin_source.name} -> {nom_pdf}")
    else:
        compteur = 1
        while chemin_pdf.exists():
            if KEEP_EXT_IN_NAME:
                nom_pdf = f"{chemin_source.name}_{compteur}.pdf"
            else:
                nom_pdf = f"{chemin_source.stem}_{compteur}.pdf"
            chemin_pdf = repertoire_dest / nom_pdf
            compteur += 1
    
    # Conversion
    print(f"📄 Conversion: {chemin_source.name} -> {nom_pdf}")
    
    reset_error_context()
    debut = time.time()
    try:
        resultat_conv = convertir_fichier_intelligent(chemin_source, chemin_pdf)
    except Exception as e_unhandled:
        log_error(f"  ❌ Exception non gérée pendant la conversion: {e_unhandled}", e_unhandled)
        resultat_conv = False
    duree = time.time() - debut
    
    if resultat_conv is True:
        taille_source = chemin_source.stat().st_size / 1024 / 1024  # MB
        taille_pdf = chemin_pdf.stat().st_size / 1024 / 1024  # MB
        print(f"  ✅ Succès en {duree:.1f}s ({taille_source:.1f}MB → {taille_pdf:.1f}MB)")
        journaliser('success', chemin_source, chemin_pdf, duree, '', error_messages=' | '.join(_LAST_ERRORS), exception=_LAST_EXCEPTION, info_messages=' | '.join(_LAST_INFOS), method_used=_LAST_METHOD_USED)
        
        if not conserver_original:
            try:
                chemin_source.unlink()
                print(f"  🗑️  Original supprimé")
            except Exception as e:
                print(f"  ⚠️  Erreur suppression: {e}")
        return 'success'

    if resultat_conv == "skipped_password":
        # Nettoyer un éventuel PDF partiel
        try:
            if chemin_pdf.exists():
                chemin_pdf.unlink()
        except Exception:
            pass
        print(f"  🔒 Ignoré (mot de passe requis): {chemin_source.name}")
        journaliser('skipped_password', chemin_source, None, duree, 'mot de passe requis / fichier protégé', error_messages=' | '.join(_LAST_ERRORS), exception=_LAST_EXCEPTION)
        return 'skipped'

    else:
        log_error(f"  ❌ Échec de conversion")
        journaliser('failed', chemin_source, chemin_pdf, duree, ('échec conversion' + (': ' + _LAST_ERRORS[0] if _LAST_ERRORS else '')), error_messages=' | '.join(_LAST_ERRORS), exception=_LAST_EXCEPTION)
        return 'failed'

def traiter_repertoire(repertoire, recursif=False, repertoire_sortie=None, 
                      conserver_original=True, extensions=None, forcer=False, journal=False):
    """Traite tous les fichiers d'un répertoire"""
    repertoire_path = Path(repertoire)
    
    if not repertoire_path.exists() or not repertoire_path.is_dir():
        print(f"Erreur: '{repertoire}' n'est pas un répertoire valide.")
        return
    
    # Extensions par défaut
    if extensions is None:
        extensions = ['.xml', '.xlsx', '.xls', '.xlsm', '.xlsb', '.docx', '.doc', '.rtf', '.odt',
                     '.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.webp',
                     '.htm', '.html', '.txt', '.log', '.msg',
                     '.ppt', '.pptx', '.pdf']
    
    repertoire_sortie_path = None
    if repertoire_sortie:
        repertoire_sortie_path = Path(repertoire_sortie)
    
    # Afficher la configuration
    afficher_configuration()

    # Journal (CSV)
    if journal:
        dossier_journal = repertoire_sortie_path if repertoire_sortie_path else repertoire_path
        init_journal(dossier_journal)
    
    # Compteurs
    fichiers_traites = 0
    conversions_reussies = 0
    fichiers_ignores = 0
    echecs = 0
    
    # Parcourir les fichiers
    pattern = '**/*' if recursif else '*'
    
    print(f"📁 Traitement: {repertoire}")
    print(f"   Mode récursif: {'Oui' if recursif else 'Non'}")
    print(f"   Extensions: {', '.join(extensions)}")
    print(f"   Conserver originaux: {'Oui' if conserver_original else 'Non'}")
    print(f"   Forcer reconversion: {'Oui' if forcer else 'Non'}")
    if UTILISER_OCR:
        print(f"   OCR activé: {MOTEUR_OCR}")
    print("-" * 60)
    
    debut_total = time.time()
    
    try:
        for fichier in repertoire_path.glob(pattern):
            if fichier.is_file() and fichier.suffix.lower() in extensions:
                fichiers_traites += 1
                
                resultat = convertir_fichier(fichier, repertoire_sortie_path, conserver_original, forcer)
                
                if resultat == 'success':
                    conversions_reussies += 1
                elif resultat == 'skipped':
                    fichiers_ignores += 1
                elif resultat == 'failed':
                    echecs += 1
    except KeyboardInterrupt:
        print("\n⛔ Interruption clavier (Ctrl+C) : arrêt propre du traitement.")
    
    duree_totale = time.time() - debut_total
    
    # Résumé
    print("\n" + "=" * 60)
    print("📊 RÉSUMÉ")
    print(f"   Durée totale: {duree_totale:.1f}s")
    print(f"   Fichiers traités: {fichiers_traites}")
    print(f"   ✅ Conversions réussies: {conversions_reussies}")
    if fichiers_ignores > 0:
        print(f"   ⏭️  Fichiers ignorés: {fichiers_ignores}")
    if echecs > 0:
        print(f"   ❌ Échecs: {echecs}")
    
    if fichiers_ignores > 0 and not forcer:
        print(f"\n💡 Utilisez --force pour reconvertir les fichiers existants")

def main():
    global UTILISER_OCR, MOTEUR_OCR, METHODE_CONVERSION, KEEP_EXT_IN_NAME
    """Fonction principale"""
    global METHODE_CONVERSION, UTILISER_OCR, MOTEUR_OCR
    
    if len(sys.argv) < 2 or '--help' in sys.argv or '-h' in sys.argv:
        print("Usage: python convertir_pdf.py <répertoire> [options]")
        print("\n📋 OPTIONS:")
        print("  -r, --recursif       : Traiter aussi les sous-répertoires")
        print("  -o, --output DIR     : Répertoire de sortie")
        print("  -d, --delete         : Supprimer les originaux après conversion")
        print("  -f, --force          : Forcer la reconversion des PDF existants")
        print("\n🔧 MÉTHODES DE CONVERSION:")
        print("  --method auto        : Détection automatique (défaut)")
        print("  --method office      : Forcer Microsoft Office")
        print("  --method libreoffice : Forcer LibreOffice")
        print("  --method reportlab   : Forcer ReportLab (basique)")
        print("\n🔤 OPTIONS OCR (pour images):")
        print("  --ocr                : Activer l'OCR pour les images")
        print("  --ocr-engine ENGINE  : Choisir le moteur OCR")
        print("                         (tesseract, easyocr, paddleocr, auto)")
        print("\n📄 FILTRES DE FORMATS:")
        print("  -x, --xml-only       : Seulement XML")
        print("  -i, --images-only    : Seulement images")
        print("  -e, --excel-only     : Seulement Excel")
        print("  -w, --word-only      : Seulement Word")
        print("  -p, --powerpoint-only: Seulement PowerPoint")
        print("\n🔍 AUTRES:")
        print("  --check              : Vérifier la configuration")
        print("  --journal            : (optionnel) Créer un journal CSV (activé par défaut)")
        print("  --no-journal         : Désactiver le journal CSV")
        print("  --log-all            : Journaliser aussi les succès/skip (par défaut: erreurs uniquement)")
        print("  --enable-reportlab-fallback : Autoriser le fallback ReportLab en mode auto")
        print("  --no-keep-ext        : Nommer en x.pdf au lieu de x.ext.pdf (par défaut: x.ext.pdf)")
        print("  -h, --help           : Afficher cette aide")
        print("\n📚 FORMATS SUPPORTÉS:")
        print("  Images     : .jpg .jpeg .png .bmp .tif .tiff .webp")
        print("  Excel      : .xlsx .xls .xlsm .xlsb")
        print("  Word       : .docx .doc .rtf .odt")
        print("  PowerPoint : .pptx .ppt")
        print("  Données    : .xml")
        print("  Web        : .htm .html")
        print("  Texte      : .txt .log")
        print("  Email      : .msg")
        print("  PDF        : .pdf (copie/skip)")
        print("\n💡 EXEMPLES:")
        print("  python convertir_pdf.py ./documents")
        print("  python convertir_pdf.py ./documents -r -o ./pdf_output")
        print("  python convertir_pdf.py ./documents --method office")
        print("  python convertir_pdf.py ./scans --images-only --ocr")
        print("  python convertir_pdf.py ./scans --ocr --ocr-engine tesseract")
        sys.exit(0)
    
    # Vérification de configuration seulement
    if '--check' in sys.argv:
        afficher_configuration()
        sys.exit(0)
    
    # Parser les arguments
    repertoire = sys.argv[1]
    recursif = '-r' in sys.argv or '--recursif' in sys.argv
    supprimer_originaux = '-d' in sys.argv or '--delete' in sys.argv
    forcer = '-f' in sys.argv or '--force' in sys.argv

    # Nom des PDFs
    global KEEP_EXT_IN_NAME
    KEEP_EXT_IN_NAME = ('--no-keep-ext' not in sys.argv)
    
    # OCR
    UTILISER_OCR = '--ocr' in sys.argv
    if '--ocr-engine' in sys.argv:
        idx = sys.argv.index('--ocr-engine')
        if idx + 1 < len(sys.argv):
            MOTEUR_OCR = sys.argv[idx + 1].lower()
            if MOTEUR_OCR not in ['auto', 'tesseract', 'easyocr', 'paddleocr']:
                print(f"⚠️  Moteur OCR inconnu: {MOTEUR_OCR}")
                MOTEUR_OCR = 'auto'
    
    # Méthode de conversion
    if '--method' in sys.argv:
        idx = sys.argv.index('--method')
        if idx + 1 < len(sys.argv):
            METHODE_CONVERSION = sys.argv[idx + 1].lower()
            if METHODE_CONVERSION not in ['auto', 'office', 'libreoffice', 'reportlab']:
                print(f"⚠️  Méthode inconnue: {METHODE_CONVERSION}")
                METHODE_CONVERSION = 'auto'
    
    # Répertoire de sortie
    repertoire_sortie = None
    if '-o' in sys.argv:
        idx = sys.argv.index('-o')
        if idx + 1 < len(sys.argv):
            repertoire_sortie = sys.argv[idx + 1]
    elif '--output' in sys.argv:
        idx = sys.argv.index('--output')
        if idx + 1 < len(sys.argv):
            repertoire_sortie = sys.argv[idx + 1]
    
    # Extensions à traiter
    extensions = None
    if '-x' in sys.argv or '--xml-only' in sys.argv:
        extensions = ['.xml']
    elif '-i' in sys.argv or '--images-only' in sys.argv:
        extensions = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.webp']
    elif '-e' in sys.argv or '--excel-only' in sys.argv:
        extensions = ['.xlsx', '.xls', '.xlsm', '.xlsb']
    elif '-w' in sys.argv or '--word-only' in sys.argv:
        extensions = ['.docx', '.doc', '.rtf', '.odt']
    elif '-p' in sys.argv or '--powerpoint-only' in sys.argv:
        extensions = ['.pptx', '.ppt']
    
    # Vérifier les dépendances minimales
    try:
        import PIL
    except ImportError:
        print("❌ Pillow n'est pas installé!")
        print("   Installation: pip install Pillow")
        sys.exit(1)
    
    # Détecter les outils disponibles
    detecter_libreoffice()
    
    # Si OCR activé, vérifier la disponibilité
    if UTILISER_OCR:
        ocr_disponible = any([
            TESSERACT_AVAILABLE and detecter_tesseract(),
            EASYOCR_AVAILABLE,
            PADDLEOCR_AVAILABLE
        ])
        
        if not ocr_disponible:
            print("⚠️  Aucun moteur OCR disponible!")
            print("   Installation recommandée:")
            print("   • Tesseract + pip install pytesseract")
            print("   • Ou: pip install easyocr")
            print("   • Ou: pip install paddleocr")
            reponse = input("\nContinuer sans OCR? (o/N): ")
            if reponse.lower() != 'o':
                sys.exit(1)
            UTILISER_OCR = False
        else:
            print(f"🔤 OCR activé (moteur: {MOTEUR_OCR})")
    
    # Journal: par défaut erreurs uniquement (option --log-all)
    global JOURNAL_ERRORS_ONLY
    JOURNAL_ERRORS_ONLY = ('--log-all' not in sys.argv)

    # Fallback ReportLab en mode auto (désactivé par défaut)
    global REPORTLAB_FALLBACK_ENABLED
    REPORTLAB_FALLBACK_ENABLED = ('--enable-reportlab-fallback' in sys.argv)

    # Traiter le répertoire
    traiter_repertoire(
        repertoire, 
        recursif=recursif,
        repertoire_sortie=repertoire_sortie,
        conserver_original=not supprimer_originaux,
        extensions=extensions,
        forcer=forcer
    ,
        journal=('--no-journal' not in sys.argv)
    )

if __name__ == "__main__":
    main()