"""
Convertisseurs de secours via ReportLab.

Fallback pour Word et Excel quand Office et LibreOffice
ne sont pas disponibles. Qualité réduite mais fonctionnel.
"""

from __future__ import annotations

import time
from pathlib import Path
from typing import TYPE_CHECKING

from .base import BaseConverter, ConversionResult, ConversionStatus

# Import conditionnel de ReportLab
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Preformatted, Table, TableStyle
    )
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Import conditionnel de python-docx
try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    Document = None  # type: ignore

# Import conditionnel de pandas/openpyxl
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    pd = None  # type: ignore

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

if TYPE_CHECKING:
    from ..config import Config
    from ..logger import ConverterLogger


class ReportLabWordConverter(BaseConverter):
    """
    Convertisseur Word de secours via ReportLab + python-docx.

    Qualité réduite mais fonctionne sans Microsoft Office.
    Supporte uniquement .docx (pas .doc).
    """

    name = "reportlab_word"
    supported_extensions = [".docx"]  # Uniquement .docx

    def is_available(self) -> bool:
        """Vérifie que ReportLab et python-docx sont installés."""
        return REPORTLAB_AVAILABLE and PYTHON_DOCX_AVAILABLE

    def convert(self, source: Path, dest: Path) -> ConversionResult:
        """Convertit un document Word en PDF via ReportLab."""
        start = time.time()

        if not self.is_available():
            missing = []
            if not REPORTLAB_AVAILABLE:
                missing.append("reportlab")
            if not PYTHON_DOCX_AVAILABLE:
                missing.append("python-docx")
            return ConversionResult(
                status=ConversionStatus.FAILED,
                source=source,
                dest=None,
                duration=time.time() - start,
                method=self.name,
                message=f"Modules manquants: {', '.join(missing)}",
            )

        try:
            self.logger.debug(f"Lecture Word avec python-docx: {source.name}")
            doc = Document(source)

            # Configuration du PDF
            pdf_doc = SimpleDocTemplate(
                str(dest),
                pagesize=A4,
                leftMargin=2.5 * cm,
                rightMargin=2.5 * cm,
                topMargin=2.5 * cm,
                bottomMargin=2.5 * cm,
            )

            styles = getSampleStyleSheet()

            # Styles personnalisés
            style_normal = ParagraphStyle(
                "WordNormal",
                parent=styles["Normal"],
                fontName="Helvetica",
                fontSize=11,
                leading=14,
                spaceAfter=6,
            )

            style_heading1 = ParagraphStyle(
                "WordHeading1",
                parent=styles["Heading1"],
                fontName="Helvetica-Bold",
                fontSize=16,
                spaceAfter=12,
                spaceBefore=12,
                textColor=colors.HexColor("#2F5597"),
            )

            style_heading2 = ParagraphStyle(
                "WordHeading2",
                parent=styles["Heading2"],
                fontName="Helvetica-Bold",
                fontSize=14,
                spaceAfter=10,
                spaceBefore=10,
                textColor=colors.HexColor("#2F5597"),
            )

            story = []

            # Titre du document
            title = source.stem.replace("_", " ")
            story.append(Paragraph(title, style_heading1))
            story.append(Spacer(1, 20))

            # Parcourir les paragraphes
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Déterminer le style
                    style_name = para.style.name if para.style else ""
                    if style_name.startswith("Heading 1"):
                        style = style_heading1
                    elif style_name.startswith("Heading 2"):
                        style = style_heading2
                    else:
                        style = style_normal

                    # Échapper les caractères spéciaux
                    text_escaped = (
                        text.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )

                    story.append(Paragraph(text_escaped, style))

            # Parcourir les tableaux
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if len(cell_text) > 100:
                            cell_text = cell_text[:97] + "..."
                        row_data.append(cell_text)
                    table_data.append(row_data)

                if table_data:
                    num_cols = len(table_data[0])
                    col_width = (A4[0] - 5 * cm) / num_cols

                    t = Table(table_data, colWidths=[col_width] * num_cols)
                    t.setStyle(TableStyle([
                        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 12))

            # Note de conversion
            note_style = ParagraphStyle(
                "Note",
                parent=styles["Italic"],
                fontSize=8,
                textColor=colors.grey,
                alignment=1,
            )
            story.append(Spacer(1, 30))
            story.append(Paragraph(
                "Converti avec ReportLab (qualité réduite)",
                note_style,
            ))

            pdf_doc.build(story)

            if dest.exists():
                self.logger.debug("Conversion ReportLab Word réussie")
                return ConversionResult(
                    status=ConversionStatus.SUCCESS,
                    source=source,
                    dest=dest,
                    duration=time.time() - start,
                    method=self.name,
                )

            return ConversionResult(
                status=ConversionStatus.FAILED,
                source=source,
                dest=None,
                duration=time.time() - start,
                method=self.name,
                message="PDF non créé",
            )

        except Exception as e:
            self.logger.error(f"Erreur ReportLab Word: {e}", exc=e)
            return ConversionResult(
                status=ConversionStatus.FAILED,
                source=source,
                dest=None,
                duration=time.time() - start,
                method=self.name,
                exception=e,
            )


class ReportLabExcelConverter(BaseConverter):
    """
    Convertisseur Excel de secours via ReportLab + pandas.

    Qualité réduite mais fonctionne sans Microsoft Office.
    """

    name = "reportlab_excel"
    supported_extensions = [".xlsx", ".xls", ".xlsm"]

    def is_available(self) -> bool:
        """Vérifie que ReportLab et pandas sont installés."""
        return REPORTLAB_AVAILABLE and PANDAS_AVAILABLE

    def convert(self, source: Path, dest: Path) -> ConversionResult:
        """Convertit un classeur Excel en PDF via ReportLab."""
        start = time.time()

        if not self.is_available():
            missing = []
            if not REPORTLAB_AVAILABLE:
                missing.append("reportlab")
            if not PANDAS_AVAILABLE:
                missing.append("pandas")
            return ConversionResult(
                status=ConversionStatus.FAILED,
                source=source,
                dest=None,
                duration=time.time() - start,
                method=self.name,
                message=f"Modules manquants: {', '.join(missing)}",
            )

        try:
            self.logger.debug(f"Lecture Excel avec pandas: {source.name}")

            # Lire toutes les feuilles
            engine = "openpyxl" if OPENPYXL_AVAILABLE else None
            excel_file = pd.ExcelFile(source, engine=engine)

            # Configuration du PDF
            pdf_doc = SimpleDocTemplate(
                str(dest),
                pagesize=A4,
                leftMargin=20,
                rightMargin=20,
                topMargin=30,
                bottomMargin=30,
            )

            styles = getSampleStyleSheet()
            story = []

            # Titre principal
            title_style = ParagraphStyle(
                "ExcelTitle",
                parent=styles["Title"],
                fontName="Helvetica",
                fontSize=14,
                spaceAfter=15,
            )
            story.append(Paragraph(source.name, title_style))
            story.append(Spacer(1, 10))

            # Traiter chaque feuille
            max_rows = 100  # Limite pour éviter les PDF trop longs

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(
                    source,
                    sheet_name=sheet_name,
                    engine=engine,
                    na_filter=True,
                )

                if df.empty:
                    continue

                # Titre de la feuille
                story.append(Paragraph(f"Feuille: {sheet_name}", styles["Heading2"]))
                story.append(Spacer(1, 10))

                # Préparer les données
                data = []

                # En-têtes
                headers = []
                for col in df.columns:
                    header = str(col)
                    if "Unnamed:" in header:
                        header = ""
                    headers.append(header)
                data.append(headers)

                # Données (limitées)
                df_display = df.head(max_rows)
                for _, row in df_display.iterrows():
                    row_data = []
                    for val in row:
                        if pd.isna(val) or str(val) == "nan":
                            cell_text = ""
                        else:
                            cell_text = str(val)
                            if len(cell_text) > 50:
                                cell_text = cell_text[:47] + "..."
                        row_data.append(cell_text)
                    data.append(row_data)

                if len(data) > 1:
                    # Calculer les largeurs de colonnes
                    num_cols = len(data[0])
                    page_width = A4[0] - 40
                    col_width = page_width / num_cols
                    col_widths = [max(30, min(150, col_width))] * num_cols

                    # Créer le tableau
                    table = Table(data, colWidths=col_widths, repeatRows=1)
                    table.setStyle(TableStyle([
                        # En-têtes
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, 0), 9),
                        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                        # Données
                        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 1), (-1, -1), 8),
                        ("ALIGN", (0, 1), (-1, -1), "LEFT"),
                        # Grille
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
                        # Padding
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]))

                    story.append(table)

                    # Note si données tronquées
                    if len(df) > max_rows:
                        info_style = ParagraphStyle(
                            "InfoStyle",
                            parent=styles["Italic"],
                            fontSize=8,
                            textColor=colors.grey,
                        )
                        story.append(Spacer(1, 5))
                        story.append(Paragraph(
                            f"* Affichage limité à {max_rows} lignes sur {len(df)}",
                            info_style,
                        ))

                story.append(Spacer(1, 20))

            pdf_doc.build(story)

            if dest.exists():
                self.logger.debug("Conversion ReportLab Excel réussie")
                return ConversionResult(
                    status=ConversionStatus.SUCCESS,
                    source=source,
                    dest=dest,
                    duration=time.time() - start,
                    method=self.name,
                )

            return ConversionResult(
                status=ConversionStatus.FAILED,
                source=source,
                dest=None,
                duration=time.time() - start,
                method=self.name,
                message="PDF non créé",
            )

        except Exception as e:
            self.logger.error(f"Erreur ReportLab Excel: {e}", exc=e)
            return ConversionResult(
                status=ConversionStatus.FAILED,
                source=source,
                dest=None,
                duration=time.time() - start,
                method=self.name,
                exception=e,
            )
